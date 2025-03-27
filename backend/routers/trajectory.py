"""
轨迹和报告生成相关的路由
注意：这个路由被挂载在/api/trajectory前缀下
所有路由的完整路径为 /api/trajectory/xxx
"""

from fastapi import APIRouter, Depends, HTTPException, Request
from sqlalchemy.orm import Session
from backend.db.database import get_db, IS_PRODUCTION
from backend.models.models import Task, Image, PersonInvolved
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
import json
from fastapi.responses import FileResponse, HTMLResponse
import shutil
import mimetypes

import logging

# 配置日志
logger = logging.getLogger(__name__)

# 创建路由实例，这将被挂载在/api/trajectory前缀下
router = APIRouter(
    tags=["trajectory"],
    responses={404: {"description": "未找到资源"}}
)

# 创建outputs目录（如果不存在）
os.makedirs('outputs', exist_ok=True)

# 创建任务子目录函数
def get_task_output_dir(task_id: str) -> str:
    """获取任务输出目录，如果不存在则创建"""
    task_dir = os.path.join('outputs', f'task_{task_id}')
    os.makedirs(task_dir, exist_ok=True)
    return task_dir

def get_case_info(case_id: str, db: Session):
    """获取案件信息，包括图片和相关人员信息"""
    try:
        # 获取任务信息
        task = db.query(Task).filter(Task.id == case_id).first()
        if not task:
            return None
            
        # 获取所有相关图片信息，按时间排序
        images = db.query(Image).filter(Image.task_id == case_id).order_by(Image.time).all()
        
        # 处理图片信息
        images_info = []
        for img in images:
            # 获取该图片相关的人员信息
            people = db.query(PersonInvolved).filter(PersonInvolved.image_id == img.id).all()
            
            image_info = {
                'time': img.time.strftime('%Y-%m-%d %H:%M:%S') if img.time else '',
                'location': img.location or '',
                'transportation': img.transportation or '',
                'description': img.description or '',
                'persons': [{
                    'name': p.name,
                    'id_number': p.id_number,
                    'hometown': p.household_registration
                } for p in people],
                'image_path': img.file_path if img.file_path else None
            }
            images_info.append(image_info)
        
        # 获取所有涉事人员信息（不重复）
        involved_persons = db.query(PersonInvolved).distinct(
            PersonInvolved.name,
            PersonInvolved.id_number,
            PersonInvolved.household_registration
        ).filter(
            PersonInvolved.image_id.in_([img.id for img in images])
        ).all()
        
        return {
            'subject': task.title,  # 使用任务标题作为主题
            'images_info': images_info,
            'involved_persons': [{
                'name': p.name,
                'id_number': p.id_number,
                'hometown': p.household_registration
            } for p in involved_persons]
        }
    except Exception as e:
        logger.error(f"获取案件信息失败: {str(e)}", exc_info=not IS_PRODUCTION)
        raise HTTPException(status_code=500, detail="获取案件信息失败")

@router.get("/excel/{case_id}")
async def generate_trajectory_excel(case_id: str, db: Session = Depends(get_db)):
    """生成指定案件的轨迹表Excel文件"""
    logger.info(f"开始为案件ID {case_id} 生成轨迹表")
    try:
        # 获取案件信息
        case_info = get_case_info(case_id, db)
        if not case_info:
            raise HTTPException(status_code=404, detail="未找到相关案件信息")
            
        # 准备Excel数据
        excel_data = []
        for img_info in case_info['images_info']:
            # 将人员信息转换为字符串
            persons_str = ', '.join([f"{p['name']}({p['id_number']})" for p in img_info['persons']]) if img_info['persons'] else ''
            
            excel_data.append({
                '时间': img_info['time'],
                '地点': img_info['location'],
                '交通方式': img_info['transportation'],
                '事件描述': img_info['description'],
                '涉事人员': persons_str
            })
        
        # 创建DataFrame
        df = pd.DataFrame(excel_data)
        
        # 生成文件名
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"关于{case_info['subject']}的轨迹表_{current_time}.xlsx"
        
        # 获取任务输出目录
        task_dir = get_task_output_dir(case_id)
        filepath = os.path.join(task_dir, filename)
        
        # 创建Excel写入器
        writer = pd.ExcelWriter(filepath, engine='openpyxl')
        
        # 写入数据
        df.to_excel(writer, sheet_name=f"关于{case_info['subject']}的轨迹表", index=False)
        
        # 调整列宽
        worksheet = writer.sheets[f"关于{case_info['subject']}的轨迹表"]
        for idx, col in enumerate(df.columns):
            max_length = max(df[col].astype(str).apply(len).max(),
                           len(col))
            worksheet.column_dimensions[chr(65 + idx)].width = max_length + 2
            
        # 保存文件
        writer.close()
        
        logger.info(f"成功生成轨迹表: {filename}")
        return {"filename": filename, "file_path": os.path.join(f'task_{case_id}', filename)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成轨迹表失败: {str(e)}", exc_info=not IS_PRODUCTION)
        raise HTTPException(status_code=500, detail="生成轨迹表失败")

@router.get("/report/{case_id}")
async def generate_trajectory_report(case_id: str, db: Session = Depends(get_db)):
    """生成指定案件的轨迹报告Word文档"""
    logger.info(f"开始为案件ID {case_id} 生成轨迹报告")
    try:
        # 获取案件信息
        case_info = get_case_info(case_id, db)
        if not case_info:
            raise HTTPException(status_code=404, detail="未找到相关案件信息")
            
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f"关于{case_info['subject']}的轨迹报告", 0)
        
        # 添加涉事人员信息
        doc.add_heading('涉事人员信息', level=1)
        if case_info['involved_persons']:
            for person in case_info['involved_persons']:
                doc.add_paragraph(
                    f"姓名：{person['name']}\n"
                    f"身份证号：{person['id_number']}\n"
                    f"户籍地：{person['hometown']}\n"
                )
        else:
            doc.add_paragraph('无涉事人员信息')
        
        # 添加轨迹信息
        doc.add_heading('轨迹信息', level=1)
        for img_info in case_info['images_info']:
            # 添加时间和地点作为小标题
            doc.add_heading(f"{img_info['time']} - {img_info['location']}", level=2)
            
            # 添加详细信息
            p = doc.add_paragraph()
            p.add_run('交通方式：').bold = True
            p.add_run(f"{img_info['transportation']}\n")
            p.add_run('事件描述：').bold = True
            p.add_run(f"{img_info['description']}\n")
            
            # 添加涉事人员信息
            if img_info['persons']:
                p.add_run('涉事人员：').bold = True
                persons_str = ', '.join([f"{p['name']}({p['id_number']})" for p in img_info['persons']])
                p.add_run(f"{persons_str}\n")
            
            # 添加图片 - 直接使用原始路径
            if img_info['image_path']:
                try:
                    # 构建完整的原始图片路径
                    original_image_path = img_info['image_path']
                    if not original_image_path.startswith('/'):
                        original_image_path = '/' + original_image_path
                    
                    # 尝试查找原始图片
                    # 1. 尝试当前路径 (带有task_id)
                    full_path = os.path.join('uploads', original_image_path.lstrip('/'))
                    
                    # 2. 尝试不带task_id的路径
                    if not os.path.exists(full_path):
                        basename = os.path.basename(original_image_path)
                        full_path = os.path.join('uploads', basename)
                    
                    # 3. 尝试用图片名直接在uploads目录下查找
                    if not os.path.exists(full_path):
                        basename = os.path.basename(original_image_path)
                        # 在uploads子目录中查找
                        for root, dirs, files in os.walk('uploads'):
                            for file in files:
                                if file == basename:
                                    full_path = os.path.join(root, file)
                                    break
                            if os.path.exists(full_path):
                                break
                    
                    if os.path.exists(full_path):
                        # 直接将原始图片添加到文档
                        doc.add_picture(full_path, width=Inches(6))
                    else:
                        logger.warning(f"找不到图片: {original_image_path}")
                        p.add_run("\n[图片未找到]\n")
                except Exception as e:
                    logger.error(f"添加图片失败: {str(e)}", exc_info=not IS_PRODUCTION)
                    p.add_run("\n[图片添加失败]\n")
            
            # 添加分隔线
            doc.add_paragraph('_' * 40)
        
        # 添加落款
        doc.add_paragraph(f"\n\n{datetime.now().strftime('%Y年%m月%d日')}")
        
        # 生成文件名和保存文件
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"关于{case_info['subject']}的轨迹报告_{current_time}.docx"
        
        # 获取任务输出目录
        task_dir = get_task_output_dir(case_id)
        filepath = os.path.join(task_dir, filename)
        doc.save(filepath)
        
        logger.info(f"成功生成轨迹报告: {filename}")
        return {"filename": filename, "file_path": os.path.join(f'task_{case_id}', filename)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成轨迹报告失败: {str(e)}", exc_info=not IS_PRODUCTION)
        raise HTTPException(status_code=500, detail="生成轨迹报告失败")

@router.get("/download/{filename}")
async def download_file(filename: str, download: bool = False):
    logger.info(f"尝试访问文件: {filename}")
    filepath = os.path.join('outputs', filename)
    if not os.path.exists(filepath):
        logger.warning(f"文件不存在: {filepath}")
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 获取文件的MIME类型
    content_type, _ = mimetypes.guess_type(filepath)
    if content_type is None:
        if filename.endswith('.xlsx'):
            content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif filename.endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            content_type = 'application/octet-stream'
    
    headers = {
        'Content-Type': content_type
    }
    
    # 解决中文文件名编码问题
    # 使用URL编码而不是latin-1编码
    import urllib.parse
    encoded_filename = urllib.parse.quote(filename)
    
    # 如果download参数为True，设置为下载模式
    if download:
        logger.info(f"开始下载文件: {filepath}")
        headers['Content-Disposition'] = f'attachment; filename="{encoded_filename}"'
    else:
        logger.info(f"开始预览文件: {filepath}")
        # 默认inline方式，浏览器会尝试预览
        headers['Content-Disposition'] = f'inline; filename="{encoded_filename}"'
    
    return FileResponse(filepath, headers=headers)

@router.get("/preview/{filename}")
async def preview_file(filename: str, request: Request):
    """预览文件，提供多种预览方式包括Microsoft Office Web Viewer和Google Docs Viewer"""
    logger.info(f"尝试预览文件: {filename}")
    filepath = os.path.join('outputs', filename)
    if not os.path.exists(filepath):
        logger.warning(f"文件不存在: {filepath}")
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 获取文件的MIME类型
    content_type, _ = mimetypes.guess_type(filepath)
    
    # 解决中文文件名编码问题
    import urllib.parse
    encoded_filename = urllib.parse.quote(filename)
    
    # 对于Office文件（Word和Excel），使用多种在线预览服务
    if filename.endswith(('.docx', '.xlsx')):
        # 获取当前主机信息，构建完整URL
        host = request.headers.get('host', 'localhost:8000')
        # 确保使用正确的协议
        protocol = "https" if request.url.scheme == "https" else "http"
        
        # 构建完整的文件URL - 必须是外部可访问的
        file_url = f"{protocol}://{host}/api/trajectory/download/{encoded_filename}"
        logger.info(f"生成外部访问URL: {file_url}")
        
        # 构建多种在线预览服务的URL
        encoded_file_url = urllib.parse.quote(file_url)
        ms_office_viewer_url = f"https://view.officeapps.live.com/op/embed.aspx?src={encoded_file_url}"
        google_docs_viewer_url = f"https://docs.google.com/viewer?url={encoded_file_url}&embedded=true"
        logger.info(f"Microsoft Office Viewer URL: {ms_office_viewer_url}")
        logger.info(f"Google Docs Viewer URL: {google_docs_viewer_url}")
        
        # 返回HTML页面，提供多种预览选项和备选方案
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>预览 {filename}</title>
            <style>
                body, html {{
                    margin: 0;
                    padding: 0;
                    height: 100%;
                    font-family: Arial, sans-serif;
                }}
                .container {{
                    display: flex;
                    flex-direction: column;
                    height: 100%;
                }}
                .toolbar {{
                    padding: 10px;
                    background-color: #f0f0f0;
                    border-bottom: 1px solid #ddd;
                    display: flex;
                    align-items: center;
                    justify-content: space-between;
                }}
                .toolbar h2 {{
                    margin: 0;
                    font-size: 16px;
                }}
                .viewer-container {{
                    flex: 1;
                    position: relative;
                }}
                iframe {{
                    width: 100%;
                    height: 100%;
                    border: none;
                    position: absolute;
                    top: 0;
                    left: 0;
                }}
                .fallback-message {{
                    display: none;
                    text-align: center;
                    padding: 20px;
                }}
                .toolbar-actions {{
                    display: flex;
                    gap: 10px;
                }}
                button {{
                    padding: 5px 10px;
                    border: 1px solid #ddd;
                    background: white;
                    border-radius: 4px;
                    cursor: pointer;
                }}
                button:hover {{
                    background: #f8f8f8;
                }}
                .active {{
                    background: #4a89dc;
                    color: white;
                    border-color: #4a89dc;
                }}
                .hidden {{
                    display: none;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="toolbar">
                    <h2>预览: {filename}</h2>
                    <div class="toolbar-actions">
                        <button id="ms-office-btn" class="active" onclick="switchViewer('ms-office')">微软预览</button>
                        <button id="google-docs-btn" onclick="switchViewer('google-docs')">谷歌预览</button>
                        <a href="{file_url}" download="{filename}" target="_blank">
                            <button>下载文件</button>
                        </a>
                    </div>
                </div>
                <div class="viewer-container">
                    <iframe id="ms-office-frame" src="{ms_office_viewer_url}"></iframe>
                    <iframe id="google-docs-frame" src="{google_docs_viewer_url}" class="hidden"></iframe>
                    <div id="fallback-message" class="fallback-message">
                        <h3>预览加载失败</h3>
                        <p>无法加载文档预览。这可能是因为：</p>
                        <ul>
                            <li>您的网络连接有问题</li>
                            <li>文件格式不支持在线预览</li>
                            <li>预览服务当前不可用</li>
                        </ul>
                        <p>您可以尝试 <a href="{file_url}" download="{filename}">下载文件</a> 查看或稍后再试。</p>
                    </div>
                </div>
            </div>
            
            <script>
                // 切换预览服务
                function switchViewer(type) {{
                    // 隐藏所有iframe
                    document.querySelectorAll('iframe').forEach(iframe => {{
                        iframe.classList.add('hidden');
                    }});
                    
                    // 移除所有按钮的活动状态
                    document.querySelectorAll('.toolbar-actions button').forEach(btn => {{
                        btn.classList.remove('active');
                    }});
                    
                    // 显示选择的预览服务
                    if (type === 'ms-office') {{
                        document.getElementById('ms-office-frame').classList.remove('hidden');
                        document.getElementById('ms-office-btn').classList.add('active');
                    }} else if (type === 'google-docs') {{
                        document.getElementById('google-docs-frame').classList.remove('hidden');
                        document.getElementById('google-docs-btn').classList.add('active');
                    }}
                }}
                
                // 监控iframe加载状态
                function checkIframeLoaded(frameId) {{
                    try {{
                        const iframe = document.getElementById(frameId);
                        if (iframe.contentDocument.body.innerHTML === '') {{
                            return false;
                        }}
                        return true;
                    }} catch(e) {{
                        // 跨域问题可能导致错误，但这通常意味着iframe已加载
                        return true;
                    }}
                }}
                
                // 检查所有预览服务的加载状态
                setTimeout(function() {{
                    const msOfficeLoaded = checkIframeLoaded('ms-office-frame');
                    const googleDocsLoaded = checkIframeLoaded('google-docs-frame');
                    
                    if (!msOfficeLoaded && !googleDocsLoaded) {{
                        document.querySelectorAll('iframe').forEach(iframe => {{
                            iframe.classList.add('hidden');
                        }});
                        document.getElementById('fallback-message').style.display = 'block';
                    }}
                }}, 15000);
            </script>
        </body>
        </html>
        """
        return HTMLResponse(content=html_content, status_code=200)
    else:
        # 对于其他类型的文件，使用常规的预览方式
        if content_type is None:
            if filename.endswith('.xlsx'):
                content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            elif filename.endswith('.docx'):
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                content_type = 'application/octet-stream'
        
        headers = {
            'Content-Type': content_type,
            'Content-Disposition': f'inline; filename="{encoded_filename}"',
            'X-Content-Type-Options': 'nosniff'
        }
        
        logger.info(f"开始预览文件: {filepath}，内容类型: {content_type}")
        return FileResponse(filepath, headers=headers)

@router.get("/download-file/{filename}")
async def force_download_file(filename: str):
    """强制下载文件"""
    logger.info(f"强制下载文件: {filename}")
    filepath = os.path.join('outputs', filename)
    if not os.path.exists(filepath):
        logger.warning(f"文件不存在: {filepath}")
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 获取文件的MIME类型
    content_type, _ = mimetypes.guess_type(filepath)
    if content_type is None:
        if filename.endswith('.xlsx'):
            content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif filename.endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            content_type = 'application/octet-stream'
    
    # 解决中文文件名编码问题
    import urllib.parse
    encoded_filename = urllib.parse.quote(filename)
    
    headers = {
        'Content-Type': content_type,
        # 设置为attachment模式强制下载
        'Content-Disposition': f'attachment; filename="{encoded_filename}"'
    }
    
    logger.info(f"开始下载文件: {filepath}")
    return FileResponse(filepath, headers=headers)

@router.get("/delete-file/{filename}")
async def delete_file(filename: str):
    """删除指定文件"""
    logger.info(f"尝试删除文件: {filename}")
    # 解码文件名，如果有URL编码
    import urllib.parse
    filename = urllib.parse.unquote(filename)
    
    # 检查文件是否存在于outputs目录
    filepath = os.path.join('outputs', filename)
    if not os.path.exists(filepath):
        logger.warning(f"要删除的文件不存在: {filepath}")
        raise HTTPException(status_code=404, detail="文件不存在")
    
    try:
        # 删除文件
        os.remove(filepath)
        logger.info(f"成功删除文件: {filepath}")
        return {"status": "success", "message": "文件删除成功"}
    except Exception as e:
        logger.error(f"删除文件失败: {str(e)}", exc_info=not IS_PRODUCTION)
        raise HTTPException(status_code=500, detail=f"删除文件失败: {str(e)}")

@router.get("/files/{task_id}")
async def get_task_files(task_id: str):
    """获取指定任务的所有轨迹相关文件"""
    logger.info(f"获取任务 {task_id} 的所有轨迹文件")
    try:
        # 获取任务输出目录
        task_dir = get_task_output_dir(task_id)
        
        # 检查目录是否存在
        if not os.path.exists(task_dir):
            logger.warning(f"任务输出目录不存在: {task_dir}")
            return []
        
        # 获取目录中的所有文件
        files = []
        for filename in os.listdir(task_dir):
            filepath = os.path.join(task_dir, filename)
            if os.path.isfile(filepath):
                # 获取文件创建时间
                created_time = os.path.getctime(filepath)
                created_datetime = datetime.fromtimestamp(created_time)
                
                # 解析文件类型
                if filename.lower().endswith('.xlsx'):
                    file_type = 'excel'
                elif filename.lower().endswith('.docx') or filename.lower().endswith('.pdf'):
                    file_type = 'report'
                else:
                    file_type = 'other'
                
                # 添加文件信息
                files.append({
                    "filename": os.path.join(f'task_{task_id}', filename),  # 包含相对路径
                    "display_name": filename,  # 显示名称
                    "type": file_type,
                    "created_at": created_datetime.isoformat()
                })
        
        # 按创建时间倒序排序
        files.sort(key=lambda x: x["created_at"], reverse=True)
        
        logger.info(f"找到 {len(files)} 个文件")
        return files
    except Exception as e:
        logger.error(f"获取任务文件失败: {str(e)}", exc_info=not IS_PRODUCTION)
        raise HTTPException(status_code=500, detail=f"获取任务文件失败: {str(e)}") 